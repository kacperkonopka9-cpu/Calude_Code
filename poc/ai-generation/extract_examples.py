#!/usr/bin/env python3
"""
Extract few-shot examples from Word documents to text files.
"""

from pathlib import Path
from docx import Document
import sys

def extract_example(docx_path: Path, output_txt: Path) -> bool:
    """
    Extract text from Word document and save as UTF-8 text file.

    Args:
        docx_path: Path to source .docx file
        output_txt: Path to output .txt file

    Returns:
        True if successful, False otherwise
    """
    try:
        print(f"📄 Opening: {docx_path.name}")
        doc = Document(docx_path)

        full_text = []

        # Extract paragraphs
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                # Detect headings (usually bold or heading styles)
                if para.style.name.startswith('Heading'):
                    # Extract heading level
                    level = para.style.name.replace('Heading', '').strip()
                    if level and level.isdigit():
                        full_text.append(f"\n{'#' * int(level)} {text}\n")
                    else:
                        full_text.append(f"\n## {text}\n")
                else:
                    full_text.append(text)

        # Extract tables (if paragraphs didn't capture content)
        if not full_text or len('\n'.join(full_text)) < 100:
            print("   ℹ️  No paragraphs found, trying tables...")
            for table in doc.tables:
                for row in table.rows:
                    # Extract unique cell texts from this row
                    row_cells = []
                    seen = set()
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        # Only add if not empty and not already seen in this row
                        if cell_text and cell_text not in seen:
                            row_cells.append(cell_text)
                            seen.add(cell_text)

                    # Join unique cells for this row
                    if row_cells:
                        # If there's a separator (label | value pattern), keep it clean
                        if len(row_cells) == 2:
                            full_text.append(f"{row_cells[0]}: {row_cells[1]}")
                        elif len(row_cells) == 1:
                            full_text.append(row_cells[0])
                        else:
                            full_text.append('\n'.join(row_cells))

        # Join with double newlines for paragraphs
        output = '\n\n'.join(full_text)

        # Clean up excessive newlines
        while '\n\n\n' in output:
            output = output.replace('\n\n\n', '\n\n')

        # Write to file
        output_txt.parent.mkdir(parents=True, exist_ok=True)
        with open(output_txt, 'w', encoding='utf-8') as f:
            f.write(output)

        # Verify output
        chars = len(output)
        words = len(output.split())
        kb = chars / 1024

        print(f"✅ Extracted: {output_txt.name}")
        print(f"   - {chars:,} characters ({kb:.1f} KB)")
        print(f"   - {words:,} words")

        # Check if size is reasonable (4-6 KB expected)
        if kb < 3:
            print(f"   ⚠️  Warning: File smaller than expected (expected 4-6 KB)")
        elif kb > 10:
            print(f"   ⚠️  Warning: File larger than expected (expected 4-6 KB)")

        return True

    except Exception as e:
        print(f"❌ Failed to extract {docx_path.name}: {e}")
        return False

def main():
    """Extract all three example files."""

    # Define paths
    base_path = Path('/workspaces/Calude_Code/docs/project-cards/examples')
    output_path = Path('/workspaces/Calude_Code/poc/ai-generation/prompts/examples')

    # Define examples to extract
    examples = [
        ('11. Karta projektu_BR ERP.docx', 'example1-erp-system.txt'),
        ('10. Karta projektu_BR logistyka Genetix.docx', 'example2-logistics-genetix.txt'),
        ('10. Karta projektu_BR HAUTAU.docx', 'example3-hautau-windows.txt'),
    ]

    print("=" * 60)
    print("  Few-Shot Example Extraction")
    print("=" * 60)
    print()

    success_count = 0

    for docx_file, txt_file in examples:
        docx_path = base_path / docx_file
        txt_path = output_path / txt_file

        if not docx_path.exists():
            print(f"❌ Source file not found: {docx_file}")
            continue

        if extract_example(docx_path, txt_path):
            success_count += 1
        print()

    # Summary
    print("=" * 60)
    if success_count == len(examples):
        print(f"✅ SUCCESS: All {success_count} examples extracted!")
        print()
        print("Next steps:")
        print("1. Verify files: ls -lh poc/ai-generation/prompts/examples/*.txt")
        print("2. Test POC: python poc/ai-generation/src/generate.py --test-ids TC-001")
        return 0
    else:
        print(f"⚠️  PARTIAL: {success_count}/{len(examples)} examples extracted")
        print()
        print("Manual extraction may be needed for failed files.")
        print("See: /poc/EXTRACT-EXAMPLES.md")
        return 1

if __name__ == '__main__':
    sys.exit(main())
